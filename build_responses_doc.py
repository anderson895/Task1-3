"""
Builds the collated textual-responses document (.docx) for Tasks 1-3.
Run after all task scripts have produced their figures.
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = os.path.dirname(os.path.abspath(__file__))
doc = Document()

# ----------------------------------------------------------------- styles
BLACK = RGBColor(0, 0, 0)

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.font.color.rgb = BLACK

# Plain black, simple formatting for all heading levels (override the
# default coloured Word theme styles).
for _name in ("Title", "Heading 1", "Heading 2", "Heading 3"):
    _st = doc.styles[_name]
    _st.font.color.rgb = BLACK
    _st.font.name = "Calibri"


def h1(t): doc.add_heading(t, level=1)
def h2(t): doc.add_heading(t, level=2)
def h3(t): doc.add_heading(t, level=3)
def p(t):
    par = doc.add_paragraph(t)
    return par
def bullet(t): doc.add_paragraph(t, style="List Bullet")
def num(t): doc.add_paragraph(t, style="List Number")


def add_fig(relpath, width=6.0, caption=None):
    fp = os.path.join(ROOT, relpath)
    if os.path.exists(fp):
        doc.add_picture(fp, width=Inches(width))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            c = doc.add_paragraph(caption)
            c.alignment = WD_ALIGN_PARAGRAPH.CENTER
            c.runs[0].italic = True
            c.runs[0].font.size = Pt(9)


# ============================================================ TITLE PAGE
title = doc.add_heading("Data Science Assignment — Tasks 1–3", level=0)
sub = doc.add_paragraph("Clustering, Dimensionality Reduction, Feature Selection, "
                        "Time-Series Forecasting & Sentiment Analysis")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].italic = True
meta = doc.add_paragraph("Author: Allen Mendiola\nTools: Python, scikit-learn, "
                         "statsmodels, Matplotlib/Seaborn, TextBlob/VADER")
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph(
    "This document collates the written responses, results and interpretations "
    "for all parts. Each part has an accompanying Python script that reproduces "
    "the figures and numbers cited here.")
doc.add_page_break()

# ====================================================================
# TASK 1 - PART 1 : K-MEANS
# ====================================================================
h1("Task 1 — Part 1: Customer Segmentation with K-Means (Mall_Customers)")
p("Script: task1/part1/task1_part1_kmeans.py")
h3("Method")
bullet("Features used: Age, Annual Income (k$), Spending Score (1–100), standardised "
       "with StandardScaler. CustomerID is an identifier; Gender is kept only for profiling.")
bullet("Elbow method (inertia) and Silhouette analysis were run for k = 2..10.")
bullet("The inertia curve bends clearly around k = 5 (elbow). The silhouette score is "
       "high and stable from k = 5 (0.417) onward, peaking marginally at k = 6 (0.428). "
       "k = 5 was chosen as it sits at the elbow and yields the cleanest, most "
       "interpretable, well-separated segments.")
add_fig("task1/part1/fig1_elbow_silhouette.png", 6.3,
        "Figure 1.1.1 — Elbow method and silhouette analysis.")
add_fig("task1/part1/fig3_income_vs_spending.png", 5.2,
        "Figure 1.1.2 — Final 5 segments: Annual Income vs Spending Score (centroids in black).")

h3("Cluster profiles & business insights (k = 5)")
bullet("Cluster 2 — Premium / Target (n=40): high income (~86k), high spending (~82), "
       "avg age 33. Best ROI for loyalty programmes, premium products and personalised offers.")
bullet("Cluster 3 — Wealthy but frugal (n=39): high income (~86k), low spending (~19). "
       "Untapped potential — investigate why they don't spend; targeted incentives could convert them.")
bullet("Cluster 1 — Young enthusiasts (n=54): young (~25), moderate income (~41k), high "
       "spending (~62). Trend-driven, responsive to promotions and social campaigns.")
bullet("Cluster 4 — Mainstream / Standard (n=47): older (~56), mid income/spend. Stable "
       "core customers; retain with value bundles.")
bullet("Cluster 0 — Low value (n=20): low income (~27k), low spending (~18). Low-cost "
       "engagement only (discounts, essentials).")

h3("Question 1 — Why is choosing the number of clusters crucial? Elbow vs Silhouette.")
p("The number of clusters defines the segmentation itself. Too few clusters merge distinct "
  "customer groups (e.g. premium spenders blended with frugal high-earners), so marketing is "
  "mistargeted; too many clusters fragment the base into tiny, non-actionable groups and overfit "
  "noise. The Elbow method plots within-cluster SSE (inertia) against k and looks for the 'elbow' "
  "where adding clusters stops yielding large gains — it is fast and intuitive but subjective when "
  "the bend is gradual. Silhouette analysis measures how well each point fits its own cluster versus "
  "the nearest other cluster (range −1..1); it gives a quantitative, comparable score and also "
  "reveals mis-assigned points. Prefer the Elbow for a quick first estimate; prefer Silhouette when "
  "the elbow is ambiguous or when you need an objective, defensible choice. Using both together "
  "(as done here) is best practice.")

h3("Question 2 — Key limitations of K-Means and how to address them.")
bullet("Poor initialisation: random seeds can converge to bad local optima. Mitigation: k-means++ "
       "initialisation (default here) and n_init multiple restarts.")
bullet("Sensitivity to outliers: means are pulled by extreme values. Mitigation: remove/cap outliers, "
       "or use K-Medoids / robust scaling.")
bullet("Assumes spherical, equal-size clusters (Euclidean): fails on elongated or density-varying "
       "shapes. Mitigation: Gaussian Mixture Models, DBSCAN or spectral clustering.")
bullet("Requires k in advance and needs feature scaling: address with Elbow/Silhouette and "
       "standardisation (both done here).")
doc.add_page_break()

# ====================================================================
# TASK 1 - PART 2 : PCA
# ====================================================================
h1("Task 1 — Part 2: PCA on the Wine Quality Dataset (winequality-red)")
p("Script: task1/part2/task1_part2_pca.py")
h3("Preprocessing & results")
bullet("No missing values. 11 physicochemical features were standardised (mean 0, std 1); the "
       "target `quality` was excluded from PCA.")
bullet("Minimum components to retain ≥90% variance: 7 (cumulative variance reaches 90.83% at PC7). "
       "PC1 explains 28.2% and PC2 17.5% (45.7% together).")
add_fig("task1/part2/fig1_explained_variance.png", 6.3,
        "Figure 1.2.1 — Scree plot and cumulative explained variance (90% line at 7 components).")
add_fig("task1/part2/fig2_pca_scatter.png", 5.0,
        "Figure 1.2.2 — Wines in PC1–PC2 space coloured by quality category.")

h3("Interpretation / reflection")
bullet("Components for 90% variance: 7 of 11 — modest reduction, indicating the variance is spread "
       "across many features rather than concentrated in a few.")
bullet("2D plot observation: the three quality categories (Low 3–5, Medium 6, High 7–8) overlap "
       "heavily — there are NO clean, separable clusters by quality in the first two PCs. Quality is "
       "only weakly aligned with the directions of greatest variance.")
bullet("Feature correlations/redundancy: PC1 is an 'acidity' axis (high positive loadings on fixed "
       "acidity, citric acid, density; strong negative on pH). PC2 is a 'sulphur dioxide' axis "
       "(free + total SO₂). These groupings reveal correlated/redundant feature blocks.")

h3("Classification: original features vs reduced components")
p("A LogisticRegression classifier (3 quality classes, 5-fold CV accuracy) was trained on the "
  "original features versus the PCA-reduced components:")
bullet("Original 11 standardised features: 0.612")
bullet("PCA, 7 components (≥90% variance): 0.607 — essentially the same accuracy with 4 fewer "
       "dimensions, showing PCA preserves almost all predictive signal while reducing dimensionality.")
bullet("PCA, 2 components: 0.538 — a clear drop, confirming the first two PCs alone are insufficient "
       "(consistent with the overlapping 2D scatter).")
add_fig("task1/part2/fig4_classification_compare.png", 5.0,
        "Figure 1.2.3 — Classifier accuracy: original vs PCA-reduced features.")

h3("Question 1 — How does PCA identify redundant/less-informative features?")
p("PCA rotates the data onto orthogonal axes ordered by variance. Features that are strongly "
  "correlated load together on the same component, so their shared information is captured once — "
  "exposing redundancy (here acidity-related features cluster on PC1, SO₂ features on PC2). "
  "Components with very small explained variance (the last PCs) correspond to directions where the "
  "data barely changes; features dominating only those directions carry little discriminative "
  "information and can often be dropped, reducing dimensionality without losing much signal.")

h3("Question 2 — Why standardise before PCA, and what happens if we skip it?")
p("PCA maximises variance, and variance is scale-dependent. Without standardisation a feature with a "
  "large numeric range (e.g. total sulfur dioxide, tens–hundreds) would dominate the first components "
  "purely because of its units, while small-scale features (e.g. chlorides or density near 1) are "
  "ignored — the result reflects measurement units, not true structure. Standardising to mean 0 / "
  "variance 1 puts all features on equal footing so PCA captures genuine correlation structure.")
doc.add_page_break()

# ====================================================================
# TASK 1 - PART 3 : TIME SERIES
# ====================================================================
h1("Task 1 — Part 3: Time-Series Decomposition & Forecasting")
p("Script: task1/part3/task1_part3_timeseries.py")
h3("1. Data generation & visualisation")
bullet("Synthetic daily series of 425 points = 365 (train) + 60 (test): linear upward trend "
       "(0.10/day) + weekly 7-day sine seasonality (amplitude 10) + Gaussian noise (mean 0, std 3).")
bullet("The upward trend is clearly visible; the weekly pattern is strong and regular (amplitude 10 "
       "vs noise std 3); the noise is comparatively low, so the structure dominates.")
add_fig("task1/part3/fig1_timeseries.png", 6.3, "Figure 1.3.1 — Generated time series with train/test split.")

h3("2. Decomposition (period = 7)")
bullet("Both additive and multiplicative seasonal_decompose were applied. Because the series is "
       "built additively, the additive model captures seasonality best: its residual mean ≈ 0.009 "
       "(centred on zero). The multiplicative residuals centre on ≈ 1.0 by construction but show "
       "slightly more structure, as expected for a non-multiplicative process.")
add_fig("task1/part3/fig2_decompose_additive.png", 5.6,
        "Figure 1.3.2 — Additive decomposition (observed, trend, seasonal, residual).")

h3("3–4. Forecasting (60 days) & evaluation")
p("Three methods forecast the 60-day horizon. Accuracy on the test set:")
bullet("7-day Moving Average — MAE 7.86, MSE 80.91, MAPE 8.57%")
bullet("Simple Exponential Smoothing (SES) — MAE 7.62, MSE 74.72, MAPE 8.42%  (best)")
bullet("ARIMA(1,1,1) — MAE 8.64, MSE 106.68, MAPE 9.19%")
add_fig("task1/part3/fig3_forecasts.png", 6.3, "Figure 1.3.3 — Actual vs predicted (3 methods).")
p("SES gives the lowest errors here, but all three under-perform on the seasonal structure: none "
  "of them model the weekly cycle, so they essentially track the trend/level and miss the peaks and "
  "troughs. This is expected and motivates seasonal models (e.g. SARIMA / Holt-Winters).")

h3("5. Interpretation & reflection")
bullet("Why does SES flatten after training? SES forecasts a constant equal to the last smoothed "
       "level — it has no trend or seasonal term, so every future point is the same value (a flat line).")
bullet("What do decomposition residuals reveal? With residuals centred on ~0 (additive) and no "
       "remaining pattern, the trend + seasonal components captured the signal well; leftover structure "
       "would indicate an inadequate model or wrong period.")
bullet("ARIMA vs Moving Average on seasonality: plain ARIMA(1,1,1) and a Moving Average both ignore "
       "seasonality — ARIMA models autocorrelation/trend via differencing and AR/MA terms but needs a "
       "seasonal extension (SARIMA) to handle the weekly cycle; the MA simply smooths recent values.")
bullet("Real-world uses: Moving Average — quick smoothing/denoising of noisy KPIs; SES — short-horizon "
       "demand with no trend/season; ARIMA — trended series with autocorrelation (sales, finance).")
p("Discussion — Why decompose before forecasting? It exposes trend, seasonality and noise separately, "
  "so you choose a model that matches the structure (e.g. a seasonal model when a strong weekly cycle "
  "exists) and can de-seasonalise data. ARIMA's limitation in seasonal data: non-seasonal ARIMA cannot "
  "represent repeating cycles; it needs SARIMA (seasonal orders) or external seasonal features, otherwise "
  "it systematically misses periodic peaks/troughs — exactly what we observed.")
doc.add_page_break()

# ====================================================================
# TASK 2 - PART 1 : HIERARCHICAL
# ====================================================================
h1("Task 2 — Part 1: Hierarchical Clustering (Wholesale customers)")
p("Script: task2/part1/task2_part1_hierarchical.py")
h3("Method & linkage comparison")
bullet("Channel and Region were dropped (categorical IDs). The six spend features were log1p-"
       "transformed (heavy right skew) then standardised.")
bullet("Dendrograms were built for single, complete, average and ward linkage. Single and average "
       "linkage 'chain' — they peel off 1–2 outliers and leave one giant cluster (sizes 438/1/1), so "
       "their high silhouette (~0.5) is misleading. Complete (402/33/5) and especially Ward "
       "(262/53/125) give balanced, interpretable clusters; Ward was chosen with k = 3.")
add_fig("task2/part1/fig1_dendrograms.png", 6.3,
        "Figure 2.1.1 — Dendrograms for four linkage methods.")
add_fig("task2/part1/fig3_cluster_heatmap.png", 5.6,
        "Figure 2.1.2 — Mean annual spend per cluster (Ward, k=3).")

h3("Business insights (Ward, k = 3)")
bullet("Cluster 0 (n=262) — 'Fresh-focused': high Fresh spend, low Grocery/Detergents. Likely "
       "restaurants, cafés and HoReCa buyers. Push fresh-produce deals and reliable cold-chain delivery.")
bullet("Cluster 2 (n=125) — 'Retail / Grocery': high Milk, Grocery, Detergents_Paper and Delicatessen. "
       "Supermarkets / large retailers. Offer bulk pricing and category bundles.")
bullet("Cluster 1 (n=53) — 'Small grocery': moderate Grocery/Detergents, low Fresh. Smaller stores; "
       "mid-tier bundles and growth incentives.")

h3("Question 1 — Why choose the right number of clusters? How do dendrograms help? Linkage effect.")
p("As in K-Means, the cluster count determines whether segments are meaningful and actionable. A "
  "dendrogram visualises the full merge hierarchy and the distance at which clusters join; cutting it "
  "where there is a large vertical gap (a long un-merged branch) suggests a natural number of clusters "
  "without committing to k in advance. The linkage method defines how inter-cluster distance is measured "
  "and strongly shapes the result: single linkage (nearest points) tends to chain and produce long, "
  "stringy clusters sensitive to outliers; complete linkage (farthest points) yields compact, similar-"
  "sized clusters; average linkage is a middle ground; Ward minimises within-cluster variance and gives "
  "balanced, spherical clusters — usually the most interpretable for customer segmentation, which is why "
  "single/average chained here while Ward did not.")

h3("Question 2 — Limitations of hierarchical clustering vs K-Means; mitigation for wholesale data.")
bullet("Computational cost: O(n²) memory/time — does not scale to very large customer bases (K-Means "
       "is much faster). Mitigate with sampling or Ward on a representative subset.")
bullet("No reassignment: early merges are permanent, so mistakes propagate (K-Means iteratively "
       "reassigns). Mitigate by comparing linkages and validating with silhouette.")
bullet("Sensitivity to outliers/scale and skew: heavy-tailed spend distorts distances. Mitigate with "
       "log transform + standardisation (done here) and outlier capping.")
bullet("Shared with K-Means: results depend on distance metric and scaling; both struggle with "
       "arbitrarily shaped clusters (use DBSCAN there).")
doc.add_page_break()

# ====================================================================
# TASK 2 - PART 2 : CHURN FEATURE SELECTION
# ====================================================================
h1("Task 2 — Part 2: Preprocessing & Feature Selection (Telco Churn)")
p("Script: task2/part2/task2_part2_churn.py")
h3("1. Preprocessing & scaling")
bullet("Dropped customerID. TotalCharges had 11 blank entries → coerced to numeric and imputed with "
       "the median. Categorical columns one-hot encoded → 40 features. Target: Churn (Yes=1).")
bullet("Scaler boxplots: RobustScaler handles outliers best (it uses median/IQR, so extreme "
       "MonthlyCharges/TotalCharges values are compressed); StandardScaler centres values around zero "
       "(mean 0, unit variance); MinMaxScaler squeezes into [0,1] but is the most distorted by outliers.")
add_fig("task2/part2/fig1_scaler_boxplots.png", 6.3,
        "Figure 2.2.1 — Numeric features under MinMax / Standard / Robust scaling.")

h3("2. Variance Threshold")
bullet("VarianceThreshold(0.01) removed 0 features. After encoding, every column has variance above "
       "0.01 — there are no near-constant dummies (no category is rarer than ~1%), so nothing is "
       "discarded. The filter is still informative: it confirms there is no dead/constant feature to drop.")

h3("3. SelectKBest (chi²) vs RFE")
bullet("SelectKBest(chi², k=10) selected mostly service/contract dummies: Fiber-optic internet, "
       "OnlineSecurity_No, TechSupport_No, Contract_Month-to-month, Contract_Two-year, "
       "Electronic-check payment, and several 'No internet service' indicators.")
bullet("RFE(LogisticRegression, 10) selected the numeric drivers tenure, MonthlyCharges, TotalCharges "
       "plus Fiber-optic, Contract terms and service indicators.")
bullet("Overlap = 5 (Contract_Month-to-month, Contract_Two-year, InternetService_Fiber-optic, "
       "DeviceProtection_No-internet-service, TechSupport_No-internet-service).")
bullet("Why they differ: chi² is a univariate filter scoring each feature's association with the "
       "target independently, so it favours strong single dummies and (being non-negative-only) ignores "
       "the continuous tenure/charges directly; RFE is a multivariate wrapper that ranks features by their "
       "contribution within a logistic model, so it keeps the numeric predictors and accounts for "
       "interactions/redundancy. Filters are faster; wrappers are model-aware but costlier.")

h3("4. LASSO for feature selection")
bullet("LassoCV (α ≈ 0.00145 by 5-fold CV) kept 21 of 40 features with non-zero coefficients, "
       "shrinking the rest to exactly zero. Strongest signals: TotalCharges/tenure (negative → longer "
       "tenure lowers churn), Fiber-optic internet, Month-to-month contract, Electronic-check payment, "
       "no TechSupport/OnlineSecurity (all positive → higher churn).")
add_fig("task2/part2/fig2_lasso_coefficients.png", 5.6,
        "Figure 2.2.2 — LASSO non-zero coefficients.")
bullet("How LASSO handles correlated/redundant features: its L1 penalty tends to pick one feature from "
       "a correlated group and zero out the others, producing a sparse, less-redundant model — though the "
       "particular one chosen can be somewhat arbitrary among correlated features.")

h3("5. Final pipeline & evaluation")
bullet("Pipeline: median imputation + StandardScaler (numeric) and most-frequent imputation + OneHot "
       "(categorical) → SelectFromModel(LassoCV) → LogisticRegression. LASSO selected 19 features.")
bullet("Test metrics — Accuracy 0.805, ROC-AUC 0.843, RMSE 0.371, MAE 0.274. Strong, well-calibrated "
       "churn classifier from a compact feature set.")

h3("Discussion questions")
p("1. Why encode AND scale before LASSO/PCA with mixed data types? LASSO and PCA are numeric and "
  "scale-sensitive: categorical variables must be encoded (e.g. one-hot) to become numeric, and all "
  "features must be scaled so the L1 penalty / variance maximisation treats them comparably. Skipping "
  "scaling lets large-range features dominate and distorts which coefficients are shrunk or which "
  "components are formed.")
p("2. How do embedded methods (LASSO) differ from filter and wrapper methods? Filter methods (e.g. "
  "chi², variance threshold) score features independently of any model — fast but ignore feature "
  "interactions. Wrapper methods (e.g. RFE) repeatedly train a model on feature subsets — accurate but "
  "expensive. Embedded methods like LASSO perform selection inside model training via a penalty term, "
  "getting model-aware selection at a single-fit cost — a balance between the two.")
doc.add_page_break()

# ====================================================================
# TASK 3 - PART 1 : DBSCAN
# ====================================================================
h1("Task 3 — Part 1: DBSCAN Clustering (Wine dataset)")
p("Script: task3/part1/task3_part1_dbscan.py")
h3("Method & results")
bullet("13 physicochemical features standardised. min_samples = 10 (data has only 178 rows). The "
       "k-distance graph (k=10) gives a knee around eps ≈ 3.1, but at that radius almost all points are "
       "density-connected into one blob — the clusters have similar density and no clear density gap.")
bullet("Experimenting over eps × min_samples, the most informative configuration is eps = 2.2, "
       "min_samples = 10 → 2 clusters with silhouette 0.394 and 89 noise points.")
add_fig("task3/part1/fig1_kdistance.png", 5.2, "Figure 3.1.1 — k-distance graph for choosing eps.")
add_fig("task3/part1/fig2_dbscan_pca_tsne.png", 6.3,
        "Figure 3.1.2 — DBSCAN clusters in PCA and t-SNE space (noise marked).")

h3("Noise & insights")
bullet("Outliers/noise (label −1, 89 points) are wines that lie in low-density regions between the "
       "dense cores — chemically 'in-between' wines that don't fit either dense group. DBSCAN flagging "
       "them prevents them from distorting the cluster profiles (unlike K-Means, which forces every "
       "point into a cluster).")
bullet("Cluster 0: higher Flavanoids (~2.7) and Proline (~926) — richer, fuller-bodied profile. "
       "Cluster 1: high Malic acid (~3.8) and Color intensity (~8.1), low Flavanoids — a chemically "
       "distinct wine class. This separation is useful for quality control and wine-type identification.")
bullet("Note: this dataset is roughly convex with similar-density groups, so DBSCAN is not the ideal "
       "algorithm here (K-Means/GMM separate the three known cultivars more cleanly) — itself a useful "
       "finding about when DBSCAN does and doesn't shine.")

h3("Question 1 — Why is DBSCAN suitable for wine-type data (arbitrary shapes + noise)?")
p("DBSCAN groups points by density rather than distance to a centroid, so it can discover clusters of "
  "arbitrary, non-spherical shape and does not require the number of clusters in advance. Crucially it "
  "labels low-density points as noise, isolating anomalous or borderline wines instead of forcing them "
  "into a group. For complex chemical data with irregular structure and outliers this is advantageous — "
  "though, as seen above, it works best when clusters actually differ in density/are separated by sparse "
  "regions.")

h3("Question 2 — How do eps and min_samples affect the result?")
p("eps sets the neighbourhood radius and min_samples the minimum points to form a dense core. Small eps "
  "(or large min_samples) makes the density requirement strict → many points become noise and clusters "
  "fragment or vanish (e.g. eps=2.0 left 139 points as noise). Large eps (or small min_samples) merges "
  "everything into one cluster and hides structure (eps≈3.1 gave a single cluster). The two must be tuned "
  "together — typically min_samples ≥ dim+1 and eps from the knee of the k-distance graph — trading off "
  "the number of clusters against how many points are treated as noise.")
doc.add_page_break()

# ====================================================================
# TASK 3 - PART 2 : SENTIMENT
# ====================================================================
h1("Task 3 — Part 2: Rule-Based Sentiment Analysis (Tech Tweets)")
p("Script: task3/part2/task3_part2_sentiment.py")
h3("Method & results")
bullet("5,000 tweets (balanced: ~1,650–1,684 per class). Cleaning: lowercase; removed URLs, @mentions, "
       "the # symbol, punctuation and numbers; collapsed whitespace.")
bullet("Polarity scored with TextBlob and VADER. Labels: polarity > 0 → positive, < 0 → negative, "
       "= 0 → neutral (VADER used its standard ±0.05 compound thresholds).")
bullet("Accuracy vs the ground-truth labels: TextBlob 0.670, VADER 0.691.")
add_fig("task3/part2/fig_cm_vader.png", 4.6, "Figure 3.2.1 — VADER confusion matrix.")
add_fig("task3/part2/fig_cm_textblob.png", 4.6, "Figure 3.2.2 — TextBlob confusion matrix.")

h3("Interpretation & reflection")
bullet("How well did the rule-based method perform? Moderately — ~67–69% accuracy. Both tools handle "
       "clearly positive and negative tweets well (positive recall ≈ 1.0, negative recall 0.7–0.8) but "
       "fail badly on NEUTRAL tweets (TextBlob neutral recall 0.21, VADER 0.34): they over-predict "
       "positive, pushing many neutral tweets into the positive class.")
bullet("Examples of wrong predictions: 'Still exploring Apple. Too early to judge.' (neutral) → predicted "
       "positive; 'Amazon update broke everything. Please fix it!' (negative) → predicted neutral (idiomatic "
       "negativity missed); 'Not good, not bad. Apple is just average.' (neutral) → mis-scored due to the "
       "word 'good/bad' cancelling imperfectly.")
bullet("Common issues: sarcasm and irony, negation, mixed sentiment in one tweet, and domain/idiomatic "
       "phrasing ('broke everything') that lexicons don't weight correctly. Neutral is the hardest class "
       "because rule-based polarity rarely lands exactly at zero.")
bullet("When prefer rule-based over ML? When there is no labelled training data, when you need a fast, "
       "transparent, explainable baseline, for short social text (VADER is tuned for it), or limited "
       "compute. ML/transformer models win when labelled data exists and higher accuracy / context "
       "handling (sarcasm, negation) is required.")

h3("Discussion questions")
p("How do TextBlob/VADER work? They are lexicon-and-rule based: each word has a pre-assigned sentiment "
  "score in a dictionary; the tool aggregates these word scores (VADER also applies heuristics for "
  "negation, punctuation, capitalisation and intensifiers) into an overall polarity/compound score, which "
  "is then thresholded into positive/negative/neutral. No training data is needed.")
p("Limitations of rule-based sentiment analysis: cannot truly understand context, sarcasm, irony or "
  "complex negation; struggles with mixed sentiment and domain-specific/idiomatic language; depends on "
  "lexicon coverage (out-of-vocabulary slang scores 0); and tends to mislabel neutral text — all visible "
  "in the ~30% error rate here.")

# ---------------------------------------------------- force plain black text
for _para in doc.paragraphs:
    for _run in _para.runs:
        _run.font.color.rgb = BLACK

# ---------------------------------------------------------------- save
out = os.path.join(ROOT, "Allen_Mendiola_Data_Science_Assignment.docx")
doc.save(out)
print("Saved:", out)
